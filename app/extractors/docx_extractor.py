"""
DOCX extraction via python-docx.

DOCX has no native "page" concept in the XML (pagination is a rendering-time
property, not stored in the file) so exact page-number citation isn't
possible without a layout engine. Instead we trace by SECTION: every
Heading-style paragraph starts a new section, and all body paragraphs
until the next heading are grouped under it. This is actually a better
citation unit for SOP-style docs than a page number would be — "see
section: Audit Procedure for Production Line B" is more useful to a
user than "see page 4".

We also extract tables, since SOPs often have procedure steps or
responsibility matrices in table form.
"""

from __future__ import annotations

from pathlib import Path

import docx

from app.extractors.base import BaseExtractor, RawSegment

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}


def _table_to_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


class DOCXExtractor(BaseExtractor):
    supported_extensions = (".docx",)

    def extract(self, file_path: Path) -> list[RawSegment]:
        document = docx.Document(str(file_path))
        segments: list[RawSegment] = []

        current_heading = "Document Start"
        current_paragraphs: list[str] = []

        def flush():
            text = "\n".join(p for p in current_paragraphs if p.strip())
            if text.strip():
                segments.append(
                    RawSegment(
                        text=text,
                        source_file=file_path.name,
                        location_type="section",
                        location_value=current_heading,
                    )
                )

        for para in document.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name in HEADING_STYLES and para.text.strip():
                flush()
                current_heading = para.text.strip()
                current_paragraphs = []
            else:
                if para.text.strip():
                    current_paragraphs.append(para.text)
        flush()

        # Tables are extracted as their own segments, tagged by position
        for idx, table in enumerate(document.tables, start=1):
            table_text = _table_to_text(table)
            if table_text.strip():
                segments.append(
                    RawSegment(
                        text=table_text,
                        source_file=file_path.name,
                        location_type="section",
                        location_value=f"Table {idx}",
                        extra={"is_table": True},
                    )
                )

        return segments
